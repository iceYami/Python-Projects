"""
SISTEMA DE INSERCIÓN AUTOMATIZADA DE FIRMAS - GUÍA DE INSTALACIÓN Y USO

Este sistema permite insertar de forma automatizada firmas manuscritas simuladas, nombres completos y DNI/NIF
en documentos (.pdf, .docx, .rtf) utilizando datos extraídos desde un archivo Excel.

REQUISITOS DEL SISTEMA
-------------------------------------------------------
- Python 3.8 o superior
- Compatible con Windows, macOS y Linux
- Memoria RAM: mínimo 4 GB (recomendado 8 GB)
- Espacio libre en disco: mínimo 500 MB

INSTALACIÓN DEL SISTEMA
-------------------------------------------------------
1. Descargar o clonar este repositorio en tu equipo.
2. (Opcional pero recomendado) Crear un entorno virtual:
   En Windows:
       python -m venv env
       env\Scripts\activate
   En Linux/macOS:
       python3 -m venv env
       source env/bin/activate

3. Instalar las dependencias necesarias:
       pip install -r requirements.txt

EJECUCIÓN
-------------------------------------------------------
Para ejecutar la aplicación principal:
       python main.py

O puedes ejecutar módulos específicos según necesidad
(ej. procesamiento de Excel, inserción en PDF, etc.).

NOTAS ADICIONALES
-------------------------------------------------------
- Las firmas simuladas pueden estar almacenadas en una carpeta local (por ejemplo: /firmas/) o generarse automáticamente usando fuentes de tipo manuscrito.
- Si no se especifica carpeta de salida, los documentos firmados se exportarán por defecto en:
       - Escritorio
       - Documentos
       - Descargas
(en ese orden de prioridad, si están disponibles).
- El sistema solicitará confirmación antes de modificar y exportar documentos.
- Toda la configuración puede personalizarse desde un archivo config.json o mediante la interfaz gráfica (si está habilitada).

AVISO LEGAL
-------------------------------------------------------
Este sistema está diseñado exclusivamente para fines educativos, técnicos, de prueba o generación de documentación simulada. 
No debe utilizarse con documentos reales ni con fines de suplantación de identidad. Las firmas generadas no tienen validez legal.
"""
"""
# Dependencias para Sistema de Inserción Automatizada de Firmas
# Instalar con: pip install -r requirements.txt

# Procesamiento de datos
pandas>=1.5.0
openpyxl>=3.0.0

# Interfaz gráfica
tkinter  # Incluido en Python estándar

# Procesamiento de imágenes
Pillow>=9.0.0

# Procesamiento de documentos PDF
PyPDF2>=3.0.0
reportlab>=3.6.0

# Procesamiento de documentos Word
python-docx>=0.8.11

# Logging y utilidades estándar
pathlib  # Incluido en Python estándar
datetime  # Incluido en Python estándar
json     # Incluido en Python estándar
re       # Incluido en Python estándar
os       # Incluido en Python estándar
sys      # Incluido en Python estándar
io       # Incluido en Python estándar
logging  # Incluido en Python estándar

# Opcional para interfaz web (futuro)
# flask>=2.0.0
# flask-wtf>=1.0.0

#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Aplicación para Inserción Automatizada de Firmas en Documentos
Autor: Imegami
Versión: 1.1
"""
"""

import os
import sys
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
import json
from datetime import datetime
import re
from PIL import Image, ImageDraw, ImageFont
import io
import logging

# Importaciones para manejo de documentos
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from PyPDF2 import PdfReader, PdfWriter
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('signature_app.log'),
        logging.StreamHandler()
    ]
)

class ConfigurationManager:
    """Gestor de configuración de la aplicación"""
    
    def __init__(self):
        self.config_file = "config.json"
        self.default_config = {
            "output_path": "",
            "signature_scale_min": 0.5,
            "signature_scale_max": 2.0,
            "default_font": "Arial",
            "text_color": "#000000",
            "text_size": 12,
            "signature_opacity": 0.8,
            "margin_top": 10,
            "margin_bottom": 10
        }
        self.load_config()
    
    def load_config(self):
        """Carga la configuración desde archivo"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    self.config = json.load(f)
            else:
                self.config = self.default_config.copy()
        except Exception as e:
            logging.error(f"Error cargando configuración: {e}")
            self.config = self.default_config.copy()
    
    def save_configuration(self):
        """Guarda la configuración actual"""
        self.config.set('output_path', self.output_path_var.get())
        self.config.set('text_size', self.text_size_var.get())
        self.config.set('signature_opacity', self.opacity_var.get())
        self.config.save_config()
        messagebox.showinfo("Configuración", "Configuración guardada correctamente")
    
    def process_documents(self):
        """Procesa todos los documentos con los datos cargados"""
        # Validar que tenemos datos y documentos
        if self.excel_data is None or self.excel_data.empty:
            messagebox.showerror("Error", "No hay datos Excel válidos cargados")
            return
        
        if not self.selected_documents:
            messagebox.showerror("Error", "No hay documentos seleccionados")
            return
        
        # Generar resumen
        self._update_processing_summary()
        
        # Confirmar procesamiento
        if not messagebox.askyesno("Confirmación", 
                                  "¿Proceder con el procesamiento de documentos?"):
            return
        
        # Determinar ruta de salida
        output_path = self.output_path_var.get()
        if not output_path:
            output_path = self._get_default_output_path()
        
        # Crear carpeta de salida si no existe
        os.makedirs(output_path, exist_ok=True)
        
        # Procesar documentos
        self.processing_log = []
        total_operations = len(self.selected_documents) * len(self.excel_data)
        current_operation = 0
        
        for doc_path in self.selected_documents:
            for _, person in self.excel_data.iterrows():
                current_operation += 1
                self._log_message(f"Procesando ({current_operation}/{total_operations}): "
                                f"{os.path.basename(doc_path)} - {person['nombre_completo']}")
                
                # Procesar documento
                result_path = self.doc_processor.process_document(
                    doc_path, person.to_dict(), output_path
                )
                
                if result_path:
                    log_entry = {
                        'documento_original': os.path.basename(doc_path),
                        'firmante': person['nombre_completo'],
                        'dni': person['dni'],
                        'fecha_hora': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        'archivo_salida': os.path.basename(result_path),
                        'ruta_completa': str(result_path),
                        'tipo_firma': 'generada' if 'firma' not in person or pd.isna(person['firma']) else 'imagen'
                    }
                    self.processing_log.append(log_entry)
                    self._log_message(f"✅ Completado: {os.path.basename(result_path)}")
                else:
                    self._log_message(f"❌ Error procesando: {os.path.basename(doc_path)} - {person['nombre_completo']}")
        
        # Generar log Excel
        self._generate_excel_log(output_path)
        
        # Mostrar resumen final
        successful = len(self.processing_log)
        self._log_message(f"\n🎉 PROCESAMIENTO COMPLETADO")
        self._log_message(f"Total operaciones exitosas: {successful}")
        self._log_message(f"Documentos generados en: {output_path}")
        
        messagebox.showinfo("Completado", 
                          f"Procesamiento completado.\n"
                          f"Documentos generados: {successful}\n"
                          f"Ruta: {output_path}")
    
    def _update_processing_summary(self):
        """Actualiza el resumen de procesamiento"""
        self.summary_text.delete(1.0, tk.END)
        
        if self.excel_data is not None and not self.selected_documents:
            summary = f"📊 RESUMEN DE PROCESAMIENTO\n\n"
            summary += f"Personas a firmar: {len(self.excel_data)}\n"
            summary += f"Documentos seleccionados: {len(self.selected_documents)}\n"
            summary += f"Total operaciones: {len(self.excel_data) * len(self.selected_documents)}\n\n"
            
            summary += "PERSONAS:\n"
            for _, person in self.excel_data.head(5).iterrows():
                summary += f"• {person['nombre_completo']} - {person['dni']}\n"
            
            if len(self.excel_data) > 5:
                summary += f"... y {len(self.excel_data) - 5} más\n"
            
            summary += "\nDOCUMENTOS:\n"
            for doc in self.selected_documents[:5]:
                summary += f"• {os.path.basename(doc)}\n"
            
            if len(self.selected_documents) > 5:
                summary += f"... y {len(self.selected_documents) - 5} más\n"
            
            output_path = self.output_path_var.get() or self._get_default_output_path()
            summary += f"\nRUTA DE SALIDA: {output_path}\n"
        else:
            summary = "⚠️ Faltan datos o documentos para procesar"
        
        self.summary_text.insert(tk.END, summary)
    
    def _get_default_output_path(self):
        """Obtiene ruta de salida por defecto"""
        default_paths = [
            Path.home() / "Desktop" / "Documentos_Firmados",
            Path.home() / "Documents" / "Documentos_Firmados",
            Path.home() / "Downloads" / "Documentos_Firmados"
        ]
        
        for path in default_paths:
            if path.parent.exists():
                return str(path)
        
        return str(Path.cwd() / "Documentos_Firmados")
    
    def _log_message(self, message):
        """Registra mensaje en el log"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}\n"
        self.log_text.insert(tk.END, log_entry)
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        logging.info(message)
    
    def _generate_excel_log(self, output_path):
        """Genera log de actividad en Excel"""
        if not self.processing_log:
            return
        
        try:
            log_df = pd.DataFrame(self.processing_log)
            log_file = os.path.join(output_path, "log_firmado.xlsx")
            log_df.to_excel(log_file, index=False)
            self._log_message(f"📋 Log generado: {log_file}")
        except Exception as e:
            self._log_message(f"❌ Error generando log: {e}")
    
    def _show_excel_selection_dialog(self, excel_files):
        """Muestra diálogo de selección de archivos Excel"""
        selection_window = tk.Toplevel(self.root)
        selection_window.title("Seleccionar archivo Excel")
        selection_window.geometry("500x300")
        selection_window.transient(self.root)
        selection_window.grab_set()
        
        ttk.Label(selection_window, text="Archivos Excel encontrados:").pack(pady=10)
        
        listbox = tk.Listbox(selection_window)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        for file_path in excel_files:
            listbox.insert(tk.END, f"{file_path.name} ({file_path.parent})")
        
        def on_select():
            selection = listbox.curselection()
            if selection:
                selected_file = excel_files[selection[0]]
                selection_window.destroy()
                self._process_excel_file(selected_file)
        
        ttk.Button(selection_window, text="Seleccionar", 
                  command=on_select).pack(pady=10)
    
    def run(self):
        """Ejecuta la aplicación"""
        self.root.mainloop()

class ConsoleInterface:
    """Interfaz de línea de comandos"""
    
    def __init__(self):
        self.config = ConfigurationManager()
        self.excel_processor = ExcelDataProcessor()
        self.doc_processor = DocumentProcessor(self.config)
    
    def run(self, args):
        """Ejecuta la interfaz de consola"""
        if len(args) < 3:
            self.show_help()
            return
        
        excel_file = args[1]
        documents = args[2:]
        
        # Cargar datos Excel
        print("Cargando datos Excel...")
        if not self.excel_processor.load_excel_file(excel_file):
            print("❌ Error cargando archivo Excel:")
            for error in self.excel_processor.validation_errors:
                print(f"  • {error}")
            return
        
        excel_data = self.excel_processor.get_valid_data()
        print(f"✅ {len(excel_data)} registros válidos cargados")
        
        # Configurar ruta de salida
        output_path = Path.cwd() / "Documentos_Firmados"
        output_path.mkdir(exist_ok=True)
        
        # Procesar documentos
        print(f"Procesando {len(documents)} documentos...")
        processing_log = []
        
        for doc_path in documents:
            if not os.path.exists(doc_path):
                print(f"❌ Documento no encontrado: {doc_path}")
                continue
            
            print(f"📄 Procesando: {os.path.basename(doc_path)}")
            
            for _, person in excel_data.iterrows():
                result_path = self.doc_processor.process_document(
                    doc_path, person.to_dict(), output_path
                )
                
                if result_path:
                    log_entry = {
                        'documento_original': os.path.basename(doc_path),
                        'firmante': person['nombre_completo'],
                        'dni': person['dni'],
                        'fecha_hora': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        'archivo_salida': os.path.basename(result_path),
                        'tipo_firma': 'generada'
                    }
                    processing_log.append(log_entry)
                    print(f"  ✅ {person['nombre_completo']} - {os.path.basename(result_path)}")
                else:
                    print(f"  ❌ Error con {person['nombre_completo']}")
        
        # Generar log
        if processing_log:
            log_df = pd.DataFrame(processing_log)
            log_file = output_path / "log_firmado.xlsx"
            log_df.to_excel(log_file, index=False)
            print(f"📋 Log generado: {log_file}")
        
        print(f"\n🎉 Procesamiento completado")
        print(f"Documentos generados en: {output_path}")
        print(f"Total operaciones exitosas: {len(processing_log)}")
    
    def show_help(self):
        """Muestra ayuda de la interfaz de consola"""
        help_text = """
Sistema de Inserción Automatizada de Firmas - Modo Consola

Uso: python signature_app.py <archivo_excel> <documento1> [documento2] ...

Ejemplos:
  python signature_app.py datos.xlsx documento.pdf
  python signature_app.py empleados.xlsx contrato.docx acuerdo.pdf

Formatos soportados:
  • Excel: .xlsx, .xls
  • Documentos: .pdf, .docx, .rtf

El archivo Excel debe contener las columnas:
  • dni (requerido)
  • nombre_completo o (nombre, apellido1, apellido2)
  • firma (opcional - ruta a imagen de firma)
"""
        print(help_text)

def main():
    """Función principal"""
    # Verificar dependencias
    missing_deps = []
    if not PDF_AVAILABLE:
        missing_deps.append("PyPDF2, reportlab (para PDF)")
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx (para DOCX)")
    
    if missing_deps:
        print("⚠️  Dependencias faltantes:")
        for dep in missing_deps:
            print(f"  • {dep}")
        print("\nInstalar con: pip install PyPDF2 reportlab python-docx")
        print("Continuando con funcionalidades limitadas...\n")
    
    # Determinar modo de ejecución
    if len(sys.argv) > 1:
        # Modo consola
        console = ConsoleInterface()
        console.run(sys.argv)
    else:
        # Modo gráfico
        try:
            app = SignatureApp()
            app.run()
        except Exception as e:
            print(f"Error iniciando interfaz gráfica: {e}")
            print("Usa: python signature_app.py <archivo_excel> <documentos...>")

if __name__ == "__main__":
    main()config(self):
        """Guarda la configuración actual"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
        except Exception as e:
            logging.error(f"Error guardando configuración: {e}")
    
    def get(self, key):
        return self.config.get(key, self.default_config.get(key))
    
    def set(self, key, value):
        self.config[key] = value

class ExcelDataProcessor:
    """Procesador de datos desde Excel"""
    
    def __init__(self):
        self.data = None
        self.validation_errors = []
    
    def find_excel_files(self):
        """Busca archivos Excel en rutas comunes"""
        common_paths = [
            Path.home() / "Desktop",
            Path.home() / "Downloads", 
            Path.home() / "Documents",
            Path.cwd()
        ]
        
        excel_files = []
        for path in common_paths:
            if path.exists():
                excel_files.extend(path.glob("*.xlsx"))
                excel_files.extend(path.glob("*.xls"))
        
        return excel_files
    
    def load_excel_file(self, file_path):
        """Carga y valida archivo Excel"""
        try:
            self.data = pd.read_excel(file_path)
            self.validation_errors = []
            return self.validate_data()
        except Exception as e:
            self.validation_errors.append(f"Error cargando archivo: {e}")
            return False
    
    def validate_data(self):
        """Valida la estructura y contenido de los datos"""
        if self.data is None or self.data.empty:
            self.validation_errors.append("Archivo Excel vacío")
            return False
        
        # Verificar columnas requeridas
        required_columns = ['dni']
        optional_columns = ['nombre', 'apellido1', 'apellido2', 'nombre_completo', 'firma']
        
        missing_required = [col for col in required_columns if col not in self.data.columns]
        if missing_required:
            self.validation_errors.append(f"Columnas requeridas faltantes: {missing_required}")
        
        # Verificar que existe al menos una forma de obtener el nombre completo
        name_columns = [col for col in ['nombre_completo', 'nombre'] if col in self.data.columns]
        if not name_columns:
            self.validation_errors.append("Debe existir columna 'nombre_completo' o 'nombre'")
        
        # Crear nombre completo si no existe
        if 'nombre_completo' not in self.data.columns:
            name_parts = []
            if 'nombre' in self.data.columns:
                name_parts.append(self.data['nombre'].fillna(''))
            if 'apellido1' in self.data.columns:
                name_parts.append(self.data['apellido1'].fillna(''))
            if 'apellido2' in self.data.columns:
                name_parts.append(self.data['apellido2'].fillna(''))
            
            if name_parts:
                self.data['nombre_completo'] = name_parts[0]
                for part in name_parts[1:]:
                    self.data['nombre_completo'] += ' ' + part
                self.data['nombre_completo'] = self.data['nombre_completo'].str.strip()
        
        # Validar DNI
        self.data['dni'] = self.data['dni'].astype(str).str.strip()
        invalid_dni = self.data[self.data['dni'].isin(['', 'nan', 'None'])].index.tolist()
        if invalid_dni:
            self.validation_errors.append(f"DNI vacío en filas: {invalid_dni}")
        
        # Validar nombres
        invalid_names = self.data[self.data['nombre_completo'].isin(['', 'nan', 'None'])].index.tolist()
        if invalid_names:
            self.validation_errors.append(f"Nombre vacío en filas: {invalid_names}")
        
        # Verificar duplicados
        duplicated_dni = self.data[self.data.duplicated(subset=['dni'])]['dni'].tolist()
        if duplicated_dni:
            self.validation_errors.append(f"DNI duplicados: {duplicated_dni}")
        
        return len(self.validation_errors) == 0
    
    def get_validation_report(self):
        """Genera reporte de validación"""
        report = {
            'total_records': len(self.data) if self.data is not None else 0,
            'valid_records': 0,
            'errors': self.validation_errors,
            'preview': None
        }
        
        if self.data is not None and not self.data.empty:
            valid_mask = (~self.data['dni'].isin(['', 'nan', 'None'])) & \
                        (~self.data['nombre_completo'].isin(['', 'nan', 'None']))
            report['valid_records'] = valid_mask.sum()
            report['preview'] = self.data.head().to_dict('records')
        
        return report
    
    def get_valid_data(self):
        """Retorna solo los datos válidos"""
        if self.data is None:
            return pd.DataFrame()
        
        valid_mask = (~self.data['dni'].isin(['', 'nan', 'None'])) & \
                    (~self.data['nombre_completo'].isin(['', 'nan', 'None']))
        return self.data[valid_mask].copy()

class SignatureGenerator:
    """Generador de firmas manuscritas simuladas"""
    
    def __init__(self, config_manager):
        self.config = config_manager
        self.signature_fonts = self._find_signature_fonts()
    
    def _find_signature_fonts(self):
        """Busca fuentes tipo firma disponibles"""
        font_paths = [
            "fonts/",
            "assets/fonts/",
            "/System/Library/Fonts/",
            "C:/Windows/Fonts/"
        ]
        
        signature_fonts = []
        signature_keywords = ['script', 'handwriting', 'signature', 'cursive']
        
        for path in font_paths:
            if os.path.exists(path):
                for file in os.listdir(path):
                    if file.lower().endswith(('.ttf', '.otf')):
                        if any(keyword in file.lower() for keyword in signature_keywords):
                            signature_fonts.append(os.path.join(path, file))
        
        return signature_fonts
    
    def generate_signature_image(self, name, width=300, height=100):
        """Genera una imagen de firma simulada"""
        try:
            # Crear imagen en blanco
            img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(img)
            
            # Seleccionar fuente
            font_path = None
            if self.signature_fonts:
                font_path = self.signature_fonts[0]
            
            try:
                if font_path and os.path.exists(font_path):
                    font = ImageFont.truetype(font_path, size=40)
                else:
                    font = ImageFont.load_default()
            except:
                font = ImageFont.load_default()
            
            # Calcular posición centrada
            bbox = draw.textbbox((0, 0), name, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            x = (width - text_width) // 2
            y = (height - text_height) // 2
            
            # Dibujar firma con efecto manuscrito
            color = (0, 0, 139, int(255 * self.config.get('signature_opacity')))  # Azul oscuro
            draw.text((x, y), name, font=font, fill=color)
            
            # Guardar en bytes
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            
            return img_bytes
            
        except Exception as e:
            logging.error(f"Error generando firma para {name}: {e}")
            return None
    
    def load_signature_image(self, signature_path):
        """Carga imagen de firma desde archivo"""
        try:
            if os.path.exists(signature_path):
                with open(signature_path, 'rb') as f:
                    return io.BytesIO(f.read())
            return None
        except Exception as e:
            logging.error(f"Error cargando firma {signature_path}: {e}")
            return None

class DocumentProcessor:
    """Procesador de documentos para inserción de firmas"""
    
    def __init__(self, config_manager):
        self.config = config_manager
        self.signature_gen = SignatureGenerator(config_manager)
    
    def find_documents(self):
        """Busca documentos en rutas comunes"""
        common_paths = [
            Path.home() / "Desktop",
            Path.home() / "Downloads",
            Path.home() / "Documents",
            Path.cwd()
        ]
        
        documents = []
        supported_extensions = ['.pdf', '.docx', '.rtf']
        
        for path in common_paths:
            if path.exists():
                for ext in supported_extensions:
                    documents.extend(path.glob(f"*{ext}"))
        
        return documents
    
    def detect_placeholders(self, text):
        """Detecta marcadores de posición en el texto"""
        placeholders = {
            'firma': re.findall(r'<<firma>>', text, re.IGNORECASE),
            'nombre': re.findall(r'<<nombre>>', text, re.IGNORECASE),
            'dni': re.findall(r'<<dni>>', text, re.IGNORECASE),
            'nombre_completo': re.findall(r'<<nombre_completo>>', text, re.IGNORECASE)
        }
        return placeholders
    
    def process_pdf_document(self, doc_path, person_data, output_path):
        """Procesa documento PDF"""
        if not PDF_AVAILABLE:
            logging.error("PyPDF2 y reportlab no están disponibles")
            return False
        
        try:
            # Leer PDF original
            with open(doc_path, 'rb') as file:
                reader = PdfReader(file)
                writer = PdfWriter()
                
                # Procesar cada página
                for page_num, page in enumerate(reader.pages):
                    # Extraer texto para buscar placeholders
                    text = page.extract_text()
                    placeholders = self.detect_placeholders(text)
                    
                    # Si hay placeholders o es la última página, agregar firma
                    if any(placeholders.values()) or page_num == len(reader.pages) - 1:
                        # Crear overlay con firma y datos
                        overlay_buffer = self._create_pdf_overlay(person_data)
                        if overlay_buffer:
                            overlay_reader = PdfReader(overlay_buffer)
                            overlay_page = overlay_reader.pages[0]
                            page.merge_page(overlay_page)
                    
                    writer.add_page(page)
                
                # Guardar documento firmado
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                return True
                
        except Exception as e:
            logging.error(f"Error procesando PDF {doc_path}: {e}")
            return False
    
    def _create_pdf_overlay(self, person_data):
        """Crea overlay PDF con firma y datos"""
        try:
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            
            # Posiciones para inserción (ajustables)
            signature_x, signature_y = 100, 100
            text_x, text_y = 100, 70
            
            # Insertar firma
            if 'firma' in person_data and person_data['firma']:
                signature_img = self.signature_gen.load_signature_image(person_data['firma'])
            else:
                signature_img = self.signature_gen.generate_signature_image(person_data['nombre_completo'])
            
            if signature_img:
                c.drawImage(signature_img, signature_x, signature_y, width=200, height=50)
            
            # Insertar texto
            c.setFont("Helvetica", self.config.get('text_size'))
            c.drawString(text_x, text_y, f"Nombre: {person_data['nombre_completo']}")
            c.drawString(text_x, text_y - 20, f"DNI: {person_data['dni']}")
            
            c.save()
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logging.error(f"Error creando overlay PDF: {e}")
            return None
    
    def process_docx_document(self, doc_path, person_data, output_path):
        """Procesa documento DOCX"""
        if not DOCX_AVAILABLE:
            logging.error("python-docx no está disponible")
            return False
        
        try:
            doc = Document(doc_path)
            
            # Buscar placeholders en el texto
            placeholders_found = False
            for paragraph in doc.paragraphs:
                if any(marker in paragraph.text.lower() for marker in ['<<firma>>', '<<nombre>>', '<<dni>>', '<<nombre_completo>>']):
                    placeholders_found = True
                    # Reemplazar marcadores
                    paragraph.text = paragraph.text.replace('<<nombre>>', person_data['nombre_completo'])
                    paragraph.text = paragraph.text.replace('<<nombre_completo>>', person_data['nombre_completo'])
                    paragraph.text = paragraph.text.replace('<<dni>>', person_data['dni'])
                    paragraph.text = paragraph.text.replace('<<firma>>', '')
            
            # Si no hay placeholders, agregar al final
            if not placeholders_found:
                doc.add_paragraph()
                doc.add_paragraph(f"Firmado por: {person_data['nombre_completo']}")
                doc.add_paragraph(f"DNI: {person_data['dni']}")
                doc.add_paragraph("Fecha: " + datetime.now().strftime("%d/%m/%Y"))
            
            # Guardar documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            logging.error(f"Error procesando DOCX {doc_path}: {e}")
            return False
    
    def process_document(self, doc_path, person_data, output_dir):
        """Procesa un documento según su formato"""
        doc_path = Path(doc_path)
        person_dni = person_data['dni'].replace(' ', '')
        
        # Generar nombre de salida
        output_name = f"{doc_path.stem}_{person_dni}_firmado{doc_path.suffix}"
        output_path = Path(output_dir) / output_name
        
        # Procesar según extensión
        if doc_path.suffix.lower() == '.pdf':
            success = self.process_pdf_document(doc_path, person_data, output_path)
        elif doc_path.suffix.lower() == '.docx':
            success = self.process_docx_document(doc_path, person_data, output_path)
        else:
            logging.warning(f"Formato no soportado: {doc_path.suffix}")
            return None
        
        if success:
            return output_path
        return None

class SignatureApp:
    """Aplicación principal con interfaz gráfica"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Sistema de Inserción Automatizada de Firmas")
        self.root.geometry("800x600")
        
        self.config = ConfigurationManager()
        self.excel_processor = ExcelDataProcessor()
        self.doc_processor = DocumentProcessor(self.config)
        
        self.excel_data = None
        self.selected_documents = []
        self.processing_log = []
        
        self.setup_ui()
    
    def setup_ui(self):
        """Configura la interfaz de usuario"""
        # Crear notebook para pestañas
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Pestaña 1: Datos Excel
        excel_frame = ttk.Frame(notebook)
        notebook.add(excel_frame, text="Datos Excel")
        self.setup_excel_tab(excel_frame)
        
        # Pestaña 2: Documentos
        docs_frame = ttk.Frame(notebook)
        notebook.add(docs_frame, text="Documentos")
        self.setup_documents_tab(docs_frame)
        
        # Pestaña 3: Configuración
        config_frame = ttk.Frame(notebook)
        notebook.add(config_frame, text="Configuración")
        self.setup_config_tab(config_frame)
        
        # Pestaña 4: Procesamiento
        process_frame = ttk.Frame(notebook)
        notebook.add(process_frame, text="Procesamiento")
        self.setup_processing_tab(process_frame)
    
    def setup_excel_tab(self, parent):
        """Configura la pestaña de datos Excel"""
        # Botones de carga
        ttk.Button(parent, text="Seleccionar archivo Excel", 
                  command=self.load_excel_file).pack(pady=5)
        ttk.Button(parent, text="Buscar automáticamente", 
                  command=self.auto_find_excel).pack(pady=5)
        
        # Frame para información de datos
        info_frame = ttk.LabelFrame(parent, text="Información de datos")
        info_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.excel_info_text = tk.Text(info_frame, height=10)
        scrollbar = ttk.Scrollbar(info_frame, orient=tk.VERTICAL, command=self.excel_info_text.yview)
        self.excel_info_text.configure(yscrollcommand=scrollbar.set)
        
        self.excel_info_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Estado de validación
        self.validation_label = ttk.Label(parent, text="No hay datos cargados", 
                                        foreground="red")
        self.validation_label.pack(pady=5)
    
    def setup_documents_tab(self, parent):
        """Configura la pestaña de documentos"""
        # Botones de selección
        ttk.Button(parent, text="Seleccionar documentos", 
                  command=self.select_documents).pack(pady=5)
        ttk.Button(parent, text="Buscar automáticamente", 
                  command=self.auto_find_documents).pack(pady=5)
        
        # Lista de documentos seleccionados
        list_frame = ttk.LabelFrame(parent, text="Documentos seleccionados")
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.docs_listbox = tk.Listbox(list_frame)
        docs_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, 
                                      command=self.docs_listbox.yview)
        self.docs_listbox.configure(yscrollcommand=docs_scrollbar.set)
        
        self.docs_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        docs_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Botón para eliminar seleccionados
        ttk.Button(parent, text="Eliminar seleccionado", 
                  command=self.remove_selected_document).pack(pady=5)
    
    def setup_config_tab(self, parent):
        """Configura la pestaña de configuración"""
        # Ruta de salida
        ttk.Label(parent, text="Ruta de salida:").pack(anchor=tk.W, padx=10, pady=5)
        output_frame = ttk.Frame(parent)
        output_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.output_path_var = tk.StringVar(value=self.config.get('output_path'))
        ttk.Entry(output_frame, textvariable=self.output_path_var).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="Seleccionar", 
                  command=self.select_output_path).pack(side=tk.RIGHT, padx=(5, 0))
        
        # Configuraciones de firma
        ttk.Label(parent, text="Tamaño de texto:").pack(anchor=tk.W, padx=10, pady=5)
        self.text_size_var = tk.IntVar(value=self.config.get('text_size'))
        ttk.Scale(parent, from_=8, to=24, variable=self.text_size_var, 
                 orient=tk.HORIZONTAL).pack(fill=tk.X, padx=10)
        
        ttk.Label(parent, text="Opacidad de firma:").pack(anchor=tk.W, padx=10, pady=5)
        self.opacity_var = tk.DoubleVar(value=self.config.get('signature_opacity'))
        ttk.Scale(parent, from_=0.1, to=1.0, variable=self.opacity_var, 
                 orient=tk.HORIZONTAL).pack(fill=tk.X, padx=10)
        
        # Botón guardar configuración
        ttk.Button(parent, text="Guardar configuración", 
                  command=self.save_configuration).pack(pady=20)
    
    def setup_processing_tab(self, parent):
        """Configura la pestaña de procesamiento"""
        # Resumen
        summary_frame = ttk.LabelFrame(parent, text="Resumen")
        summary_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.summary_text = tk.Text(summary_frame, height=8)
        self.summary_text.pack(fill=tk.X, padx=5, pady=5)
        
        # Botón de procesamiento
        ttk.Button(parent, text="CONFIRMAR Y EXPORTAR", 
                  command=self.process_documents,
                  style='Accent.TButton').pack(pady=20)
        
        # Log de actividad
        log_frame = ttk.LabelFrame(parent, text="Log de actividad")
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.log_text = tk.Text(log_frame, height=10)
        log_scrollbar = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_scrollbar.set)
        
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    def load_excel_file(self):
        """Carga archivo Excel seleccionado manualmente"""
        file_path = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if file_path:
            self._process_excel_file(file_path)
    
    def auto_find_excel(self):
        """Busca automáticamente archivos Excel"""
        excel_files = self.excel_processor.find_excel_files()
        if not excel_files:
            messagebox.showwarning("Aviso", "No se encontraron archivos Excel en las rutas comunes")
            return
        
        if len(excel_files) == 1:
            self._process_excel_file(excel_files[0])
        else:
            # Mostrar lista para selección
            self._show_excel_selection_dialog(excel_files)
    
    def _process_excel_file(self, file_path):
        """Procesa archivo Excel seleccionado"""
        if self.excel_processor.load_excel_file(file_path):
            self.excel_data = self.excel_processor.get_valid_data()
            self._update_excel_info()
            self.validation_label.config(text="✅ Datos válidos cargados", foreground="green")
        else:
            self._update_excel_info()
            self.validation_label.config(text="❌ Errores en validación", foreground="red")
    
    def _update_excel_info(self):
        """Actualiza la información mostrada sobre los datos Excel"""
        self.excel_info_text.delete(1.0, tk.END)
        
        report = self.excel_processor.get_validation_report()
        
        info_text = f"Total de registros: {report['total_records']}\n"
        info_text += f"Registros válidos: {report['valid_records']}\n\n"
        
        if report['errors']:
            info_text += "ERRORES ENCONTRADOS:\n"
            for error in report['errors']:
                info_text += f"• {error}\n"
            info_text += "\n"
        
        if report['preview']:
            info_text += "PREVISUALIZACIÓN DE DATOS:\n"
            for i, record in enumerate(report['preview'][:5]):
                info_text += f"Registro {i+1}:\n"
                for key, value in record.items():
                    info_text += f"  {key}: {value}\n"
                info_text += "\n"
        
        self.excel_info_text.insert(tk.END, info_text)
    
    def select_documents(self):
        """Selecciona documentos manualmente"""
        file_paths = filedialog.askopenfilenames(
            title="Seleccionar documentos",
            filetypes=[
                ("Todos los soportados", "*.pdf *.docx *.rtf"),
                ("PDF", "*.pdf"),
                ("Word", "*.docx"),
                ("RTF", "*.rtf")
            ]
        )
        if file_paths:
            self.selected_documents.extend(file_paths)
            self._update_documents_list()
    
    def auto_find_documents(self):
        """Busca documentos automáticamente"""
        documents = self.doc_processor.find_documents()
        if not documents:
            messagebox.showwarning("Aviso", "No se encontraron documentos en las rutas comunes")
            return
        
        self.selected_documents.extend([str(doc) for doc in documents])
        self._update_documents_list()
    
    def _update_documents_list(self):
        """Actualiza la lista de documentos"""
        self.docs_listbox.delete(0, tk.END)
        for doc in self.selected_documents:
            self.docs_listbox.insert(tk.END, os.path.basename(doc))
    
    def remove_selected_document(self):
        """Elimina documento seleccionado de la lista"""
        selection = self.docs_listbox.curselection()
        if selection:
            index = selection[0]
            self.selected_documents.pop(index)
            self._update_documents_list()
    
    def select_output_path(self):
        """Selecciona ruta de salida"""
        path = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if path:
            self.output_path_var.set(path)
    
    def save_
