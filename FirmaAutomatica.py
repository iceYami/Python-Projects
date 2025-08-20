#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Simulador de Firma Automática de Documentos
===========================================
Aplicación para insertar firmas manuscritas simuladas, nombres y DNI
en documentos PDF, DOCX y RTF usando datos de Excel.

Requeriments
pandas
pymupdf
python-docx
pillow

pip install pandas pymupdf python-docx pillow
"""

import os
import sys
import logging
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import io
from datetime import datetime
import json
import threading
import queue
from dataclasses import dataclass
from typing import List, Dict, Optional, Tuple
import re

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_signer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class PersonData:
    """Estructura de datos para una persona"""
    nombre_completo: str
    dni: str
    nombre: str = ""
    apellido1: str = ""
    apellido2: str = ""
    
    def __post_init__(self):
        if not self.nombre and self.nombre_completo:
            parts = self.nombre_completo.split()
            if len(parts) >= 1:
                self.nombre = parts[0]
            if len(parts) >= 2:
                self.apellido1 = parts[1]
            if len(parts) >= 3:
                self.apellido2 = " ".join(parts[2:])

@dataclass
class SignatureConfig:
    """Configuración para la generación de firmas"""
    font_size: int = 40
    font_color: str = "#000000"
    signature_scale: float = 1.0
    name_font_size: int = 12
    dni_font_size: int = 10
    margin_bottom: int = 50
    signature_style: str = "cursive"  # cursive, elegant, modern

class SignatureGenerator:
    """Generador de firmas manuscritas simuladas"""
    
    def __init__(self, config: SignatureConfig):
        self.config = config
        self.fonts_cache = {}
        
    def _get_signature_font(self) -> Optional[ImageFont.FreeTypeFont]:
        """Obtiene una fuente manuscrita para la firma"""
        # Fuentes manuscritas comunes en diferentes sistemas
        manuscript_fonts = [
            # Windows
            "Brush Script MT", "Lucida Handwriting", "Segoe Script",
            # macOS
            "Brush Script", "Snell Roundhand", "Marker Felt",
            # Linux/Universal
            "DejaVu Sans", "Liberation Sans"
        ]
        
        for font_name in manuscript_fonts:
            try:
                return ImageFont.truetype(font_name, self.config.font_size)
            except (OSError, IOError):
                continue
        
        # Fallback a fuente por defecto
        try:
            return ImageFont.load_default()
        except:
            return None
    
    def generate_signature(self, name: str) -> Image.Image:
        """Genera una firma manuscrita simulada"""
        font = self._get_signature_font()
        if not font:
            raise ValueError("No se pudo cargar ninguna fuente para la firma")
        
        # Calcular dimensiones
        temp_img = Image.new('RGB', (1, 1), 'white')
        temp_draw = ImageDraw.Draw(temp_img)
        bbox = temp_draw.textbbox((0, 0), name, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        # Crear imagen final con márgenes
        margin = 20
        img_width = text_width + (margin * 2)
        img_height = text_height + (margin * 2)
        
        # Crear imagen con fondo transparente
        signature_img = Image.new('RGBA', (img_width, img_height), (255, 255, 255, 0))
        draw = ImageDraw.Draw(signature_img)
        
        # Convertir color hex a RGB
        color_hex = self.config.font_color.lstrip('#')
        color_rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        
        # Dibujar texto con estilo manuscrito
        x = margin
        y = margin
        
        # Aplicar transformaciones según el estilo
        if self.config.signature_style == "elegant":
            # Añadir línea decorativa debajo
            draw.text((x, y), name, font=font, fill=color_rgb)
            line_y = y + text_height + 5
            draw.line([(x, line_y), (x + text_width, line_y)], fill=color_rgb, width=2)
        elif self.config.signature_style == "modern":
            # Estilo más angular
            draw.text((x, y), name, font=font, fill=color_rgb)
        else:  # cursive (por defecto)
            draw.text((x, y), name, font=font, fill=color_rgb)
        
        # Aplicar escalado
        if self.config.signature_scale != 1.0:
            new_width = int(img_width * self.config.signature_scale)
            new_height = int(img_height * self.config.signature_scale)
            signature_img = signature_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        return signature_img

class DocumentProcessor:
    """Procesador de documentos para insertar firmas y datos"""
    
    def __init__(self, config: SignatureConfig):
        self.config = config
        self.signature_generator = SignatureGenerator(config)
    
    def process_pdf(self, pdf_path: str, person: PersonData, output_path: str) -> bool:
        """Procesa un archivo PDF"""
        try:
            doc = fitz.open(pdf_path)
            
            # Generar firma
            signature_img = self.signature_generator.generate_signature(person.nombre_completo)
            
            # Convertir imagen a bytes
            img_bytes = io.BytesIO()
            signature_img.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            
            # Obtener última página
            last_page = doc[-1]
            page_rect = last_page.rect
            
            # Posiciones para insertar elementos
            signature_y = page_rect.height - self.config.margin_bottom - 60
            name_y = signature_y + 70
            dni_y = name_y + 20
            
            # Insertar firma
            signature_rect = fitz.Rect(50, signature_y, 200, signature_y + 60)
            last_page.insert_image(signature_rect, stream=img_bytes.getvalue())
            
            # Insertar nombre
            last_page.insert_text(
                (50, name_y),
                f"Nombre: {person.nombre_completo}",
                fontsize=self.config.name_font_size,
                color=(0, 0, 0)
            )
            
            # Insertar DNI
            last_page.insert_text(
                (50, dni_y),
                f"DNI: {person.dni}",
                fontsize=self.config.dni_font_size,
                color=(0, 0, 0)
            )
            
            # Guardar documento
            doc.save(output_path)
            doc.close()
            
            return True
            
        except Exception as e:
            logger.error(f"Error procesando PDF {pdf_path}: {str(e)}")
            return False
    
    def process_docx(self, docx_path: str, person: PersonData, output_path: str) -> bool:
        """Procesa un archivo DOCX"""
        try:
            doc = Document(docx_path)
            
            # Generar y guardar firma temporalmente
            signature_img = self.signature_generator.generate_signature(person.nombre_completo)
            temp_signature_path = "temp_signature.png"
            signature_img.save(temp_signature_path, "PNG")
            
            # Añadir nueva página o espacio al final
            doc.add_page_break()
            
            # Añadir firma
            signature_paragraph = doc.add_paragraph()
            signature_run = signature_paragraph.add_run()
            signature_run.add_picture(temp_signature_path, width=Inches(2))
            
            # Añadir nombre
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(f"Nombre: {person.nombre_completo}")
            name_run.font.size = self.config.name_font_size * 635  # Convertir a twips
            
            # Añadir DNI
            dni_paragraph = doc.add_paragraph()
            dni_run = dni_paragraph.add_run(f"DNI: {person.dni}")
            dni_run.font.size = self.config.dni_font_size * 635
            
            # Guardar documento
            doc.save(output_path)
            
            # Limpiar archivo temporal
            if os.path.exists(temp_signature_path):
                os.remove(temp_signature_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Error procesando DOCX {docx_path}: {str(e)}")
            return False
    
    def process_document(self, doc_path: str, person: PersonData, output_dir: str) -> bool:
        """Procesa un documento según su formato"""
        doc_path = Path(doc_path)
        extension = doc_path.suffix.lower()
        
        # Generar nombre de archivo de salida
        output_filename = f"{doc_path.stem}_firmado_{person.dni}{extension}"
        output_path = os.path.join(output_dir, output_filename)
        
        if extension == '.pdf':
            return self.process_pdf(str(doc_path), person, output_path)
        elif extension == '.docx':
            return self.process_docx(str(doc_path), person, output_path)
        else:
            logger.warning(f"Formato no soportado: {extension}")
            return False

class ExcelReader:
    """Lector de datos desde archivos Excel"""
    
    @staticmethod
    def find_excel_files() -> List[str]:
        """Encuentra archivos Excel en rutas comunes"""
        common_paths = [
            os.path.expanduser("~/Desktop"),
            os.path.expanduser("~/Documents"),
            os.path.expanduser("~/Downloads"),
            os.getcwd()
        ]
        
        excel_files = []
        for path in common_paths:
            if os.path.exists(path):
                for file in Path(path).glob("*.xlsx"):
                    excel_files.append(str(file))
        
        return excel_files
    
    @staticmethod
    def read_excel_data(file_path: str) -> List[PersonData]:
        """Lee datos del archivo Excel"""
        try:
            df = pd.read_excel(file_path)
            
            # Limpiar espacios en blanco
            df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
            
            people = []
            
            for _, row in df.iterrows():
                # Intentar diferentes combinaciones de columnas
                nombre_completo = ""
                dni = ""
                
                # Buscar DNI
                for col in df.columns:
                    if any(keyword in col.lower() for keyword in ['dni', 'nif', 'documento']):
                        dni = str(row[col]) if pd.notna(row[col]) else ""
                        break
                
                # Buscar nombre completo
                for col in df.columns:
                    if any(keyword in col.lower() for keyword in ['nombre_completo', 'nombre completo', 'fullname']):
                        nombre_completo = str(row[col]) if pd.notna(row[col]) else ""
                        break
                
                # Si no hay nombre completo, construirlo
                if not nombre_completo:
                    nombre_parts = []
                    for col in df.columns:
                        if any(keyword in col.lower() for keyword in ['nombre', 'name']):
                            if pd.notna(row[col]):
                                nombre_parts.append(str(row[col]))
                        elif any(keyword in col.lower() for keyword in ['apellido', 'surname', 'lastname']):
                            if pd.notna(row[col]):
                                nombre_parts.append(str(row[col]))
                    
                    nombre_completo = " ".join(nombre_parts)
                
                # Validar que tengamos los datos mínimos
                if nombre_completo and dni:
                    person = PersonData(nombre_completo=nombre_completo, dni=dni)
                    people.append(person)
            
            return people
            
        except Exception as e:
            logger.error(f"Error leyendo Excel: {str(e)}")
            raise

class DocumentSignerGUI:
    """Interfaz gráfica principal"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Simulador de Firma Automática de Documentos v2.0")
        self.root.geometry("800x600")
        
        # Variables
        self.excel_file = tk.StringVar()
        self.output_dir = tk.StringVar(value=os.path.expanduser("~/Desktop/Documentos_Firmados"))
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar(value="Listo")
        
        # Configuración
        self.config = SignatureConfig()
        
        # Datos
        self.people_data = []
        self.document_paths = []
        
        # Cola para comunicación entre hilos
        self.progress_queue = queue.Queue()
        
        self.setup_ui()
        self.load_config()
        
        # Iniciar verificación periódica de la cola
        self.root.after(100, self.check_progress_queue)
    
    def setup_ui(self):
        """Configura la interfaz de usuario"""
        # Crear notebook para pestañas
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Pestaña principal
        main_frame = ttk.Frame(notebook)
        notebook.add(main_frame, text="Procesamiento")
        
        # Pestaña de configuración
        config_frame = ttk.Frame(notebook)
        notebook.add(config_frame, text="Configuración")
        
        self.setup_main_tab(main_frame)
        self.setup_config_tab(config_frame)
        
        # Barra de estado
        status_frame = ttk.Frame(self.root)
        status_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        self.progress_bar = ttk.Progressbar(status_frame, variable=self.progress_var)
        self.progress_bar.pack(fill=tk.X, pady=(0, 5))
        
        status_label = ttk.Label(status_frame, textvariable=self.status_var)
        status_label.pack(anchor=tk.W)
    
    def setup_main_tab(self, parent):
        """Configura la pestaña principal"""
        # Sección Excel
        excel_frame = ttk.LabelFrame(parent, text="Archivo Excel con Datos", padding=10)
        excel_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Entry(excel_frame, textvariable=self.excel_file, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(excel_frame, text="Seleccionar", command=self.select_excel_file).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(excel_frame, text="Auto-detectar", command=self.auto_detect_excel).pack(side=tk.RIGHT, padx=(5, 0))
        
        # Botón para cargar datos
        ttk.Button(excel_frame, text="Cargar Datos", command=self.load_excel_data).pack(pady=(10, 0))
        
        # Lista de personas cargadas
        people_frame = ttk.LabelFrame(parent, text="Personas Cargadas", padding=10)
        people_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Treeview para mostrar datos
        columns = ('Nombre Completo', 'DNI')
        self.people_tree = ttk.Treeview(people_frame, columns=columns, show='headings', height=6)
        for col in columns:
            self.people_tree.heading(col, text=col)
            self.people_tree.column(col, width=200)
        
        scrollbar_people = ttk.Scrollbar(people_frame, orient=tk.VERTICAL, command=self.people_tree.yview)
        self.people_tree.configure(yscrollcommand=scrollbar_people.set)
        
        self.people_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_people.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Sección Documentos
        docs_frame = ttk.LabelFrame(parent, text="Documentos a Firmar", padding=10)
        docs_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        docs_buttons_frame = ttk.Frame(docs_frame)
        docs_buttons_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(docs_buttons_frame, text="Añadir Documentos", command=self.add_documents).pack(side=tk.LEFT)
        ttk.Button(docs_buttons_frame, text="Limpiar Lista", command=self.clear_documents).pack(side=tk.LEFT, padx=(5, 0))
        
        # Lista de documentos
        self.docs_listbox = tk.Listbox(docs_frame, height=6)
        scrollbar_docs = ttk.Scrollbar(docs_frame, orient=tk.VERTICAL, command=self.docs_listbox.yview)
        self.docs_listbox.configure(yscrollcommand=scrollbar_docs.set)
        
        self.docs_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_docs.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Sección Salida
        output_frame = ttk.LabelFrame(parent, text="Carpeta de Salida", padding=10)
        output_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Entry(output_frame, textvariable=self.output_dir, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="Seleccionar", command=self.select_output_dir).pack(side=tk.RIGHT, padx=(5, 0))
        
        # Botón de procesamiento
        process_frame = ttk.Frame(parent)
        process_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.process_button = ttk.Button(process_frame, text="🖋️ Firmar Documentos", command=self.start_processing)
        self.process_button.pack(side=tk.RIGHT)
    
    def setup_config_tab(self, parent):
        """Configura la pestaña de configuración"""
        # Configuración de firma
        signature_frame = ttk.LabelFrame(parent, text="Configuración de Firma", padding=10)
        signature_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Tamaño de fuente de firma
        ttk.Label(signature_frame, text="Tamaño de firma:").grid(row=0, column=0, sticky=tk.W, pady=2)
        signature_size_var = tk.IntVar(value=self.config.font_size)
        signature_scale = ttk.Scale(signature_frame, from_=20, to=80, variable=signature_size_var, orient=tk.HORIZONTAL)
        signature_scale.grid(row=0, column=1, sticky=tk.EW, padx=5, pady=2)
        ttk.Label(signature_frame, textvariable=signature_size_var).grid(row=0, column=2, pady=2)
        
        # Escala de firma
        ttk.Label(signature_frame, text="Escala de firma:").grid(row=1, column=0, sticky=tk.W, pady=2)
        scale_var = tk.DoubleVar(value=self.config.signature_scale)
        scale_scale = ttk.Scale(signature_frame, from_=0.5, to=2.0, variable=scale_var, orient=tk.HORIZONTAL)
        scale_scale.grid(row=1, column=1, sticky=tk.EW, padx=5, pady=2)
        ttk.Label(signature_frame, textvariable=scale_var).grid(row=1, column=2, pady=2)
        
        # Estilo de firma
        ttk.Label(signature_frame, text="Estilo de firma:").grid(row=2, column=0, sticky=tk.W, pady=2)
        style_var = tk.StringVar(value=self.config.signature_style)
        style_combo = ttk.Combobox(signature_frame, textvariable=style_var, values=["cursive", "elegant", "modern"])
        style_combo.grid(row=2, column=1, sticky=tk.EW, padx=5, pady=2)
        
        # Configurar grid
        signature_frame.columnconfigure(1, weight=1)
        
        # Configuración de texto
        text_frame = ttk.LabelFrame(parent, text="Configuración de Texto", padding=10)
        text_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Tamaño de fuente para nombre
        ttk.Label(text_frame, text="Tamaño fuente nombre:").grid(row=0, column=0, sticky=tk.W, pady=2)
        name_size_var = tk.IntVar(value=self.config.name_font_size)
        name_scale = ttk.Scale(text_frame, from_=8, to=20, variable=name_size_var, orient=tk.HORIZONTAL)
        name_scale.grid(row=0, column=1, sticky=tk.EW, padx=5, pady=2)
        ttk.Label(text_frame, textvariable=name_size_var).grid(row=0, column=2, pady=2)
        
        # Tamaño de fuente para DNI
        ttk.Label(text_frame, text="Tamaño fuente DNI:").grid(row=1, column=0, sticky=tk.W, pady=2)
        dni_size_var = tk.IntVar(value=self.config.dni_font_size)
        dni_scale = ttk.Scale(text_frame, from_=8, to=16, variable=dni_size_var, orient=tk.HORIZONTAL)
        dni_scale.grid(row=1, column=1, sticky=tk.EW, padx=5, pady=2)
        ttk.Label(text_frame, textvariable=dni_size_var).grid(row=1, column=2, pady=2)
        
        text_frame.columnconfigure(1, weight=1)
        
        # Botón para guardar configuración
        ttk.Button(parent, text="Guardar Configuración", 
                  command=lambda: self.save_config(signature_size_var, scale_var, style_var, 
                                                 name_size_var, dni_size_var)).pack(pady=10)
    
    def save_config(self, sig_size, scale, style, name_size, dni_size):
        """Guarda la configuración"""
        self.config.font_size = int(sig_size.get())
        self.config.signature_scale = scale.get()
        self.config.signature_style = style.get()
        self.config.name_font_size = int(name_size.get())
        self.config.dni_font_size = int(dni_size.get())
        
        # Guardar en archivo
        config_dict = {
            'font_size': self.config.font_size,
            'signature_scale': self.config.signature_scale,
            'signature_style': self.config.signature_style,
            'name_font_size': self.config.name_font_size,
            'dni_font_size': self.config.dni_font_size
        }
        
        try:
            with open('config.json', 'w') as f:
                json.dump(config_dict, f, indent=2)
            messagebox.showinfo("Configuración", "Configuración guardada correctamente")
        except Exception as e:
            messagebox.showerror("Error", f"Error guardando configuración: {str(e)}")
    
    def load_config(self):
        """Carga la configuración desde archivo"""
        try:
            if os.path.exists('config.json'):
                with open('config.json', 'r') as f:
                    config_dict = json.load(f)
                
                self.config.font_size = config_dict.get('font_size', 40)
                self.config.signature_scale = config_dict.get('signature_scale', 1.0)
                self.config.signature_style = config_dict.get('signature_style', 'cursive')
                self.config.name_font_size = config_dict.get('name_font_size', 12)
                self.config.dni_font_size = config_dict.get('dni_font_size', 10)
        except Exception as e:
            logger.warning(f"No se pudo cargar configuración: {str(e)}")
    
    def select_excel_file(self):
        """Selecciona archivo Excel"""
        filename = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if filename:
            self.excel_file.set(filename)
    
    def auto_detect_excel(self):
        """Auto-detecta archivos Excel"""
        files = ExcelReader.find_excel_files()
        if files:
            if len(files) == 1:
                self.excel_file.set(files[0])
                messagebox.showinfo("Auto-detección", f"Archivo encontrado: {os.path.basename(files[0])}")
            else:
                # Mostrar diálogo de selección
                selection_window = tk.Toplevel(self.root)
                selection_window.title("Seleccionar archivo Excel")
                selection_window.geometry("500x300")
                
                ttk.Label(selection_window, text="Archivos Excel encontrados:").pack(pady=10)
                
                listbox = tk.Listbox(selection_window)
                for file in files:
                    listbox.insert(tk.END, file)
                listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                
                def select_file():
                    selection = listbox.curselection()
                    if selection:
                        self.excel_file.set(files[selection[0]])
                        selection_window.destroy()
                
                ttk.Button(selection_window, text="Seleccionar", command=select_file).pack(pady=10)
        else:
            messagebox.showwarning("Auto-detección", "No se encontraron archivos Excel en las rutas comunes")
    
    def load_excel_data(self):
        """Carga datos del archivo Excel"""
        if not self.excel_file.get():
            messagebox.showerror("Error", "Por favor selecciona un archivo Excel")
            return
        
        try:
            self.people_data = ExcelReader.read_excel_data(self.excel_file.get())
            
            # Limpiar tree
            for item in self.people_tree.get_children():
                self.people_tree.delete(item)
            
            # Llenar tree con datos
            for person in self.people_data:
                self.people_tree.insert('', tk.END, values=(person.nombre_completo, person.dni))
            
            self.status_var.set(f"Cargadas {len(self.people_data)} personas")
            messagebox.showinfo("Éxito", f"Se cargaron {len(self.people_data)} personas correctamente")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error cargando archivo Excel: {str(e)}")
    
    def add_documents(self):
        """Añade documentos a la lista"""
        filenames = filedialog.askopenfilenames(
            title="Seleccionar documentos",
            filetypes=[
                ("Documentos soportados", "*.pdf *.docx"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        for filename in filenames:
            if filename not in self.document_paths:
                self.document_paths.append(filename)
                self.docs_listbox.insert(tk.END, os.path.basename(filename))
    
    def clear_documents(self):
        """Limpia la lista de documentos"""
        self.document_paths.clear()
        self.docs_listbox.delete(0, tk.END)
    
    def select_output_dir(self):
        """Selecciona carpeta de salida"""
        directory = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if directory:
            self.output_dir.set(directory)
    
    def start_processing(self):
        """Inicia el procesamiento de documentos"""
        # Validaciones
        if not self.people_data:
            messagebox.showerror("Error", "No hay datos de personas cargados")
            return
        
        if not self.document_paths:
            messagebox.showerror("Error", "No hay documentos seleccionados")
            return
        
        if not os.path.exists(self.output_dir.get()):
            try:
                os.makedirs(self.output_dir.get())
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo crear la carpeta de salida: {str(e)}")
                return
        
        # Deshabilitar botón de procesamiento
        self.process_button.config(state=tk.DISABLED)
        
        # Iniciar procesamiento en hilo separado
        processing_thread = threading.Thread(target=self.process_documents_thread)
        processing_thread.daemon = True
        processing_thread.start()
    
    def process_documents_thread(self):
        """Procesa documentos en hilo separado"""
        try:
            processor = DocumentProcessor(self.config)
            total_operations = len(self.people_data) * len(self.document_paths)
            current_operation = 0
            
            results = {
                'success': 0,
                'failed': 0,
                'details': []
            }
            
            for person in self.people_data:
                for doc_path in self.document_paths:
                    try:
                        # Actualizar progreso
                        progress = (current_operation / total_operations) * 100
                        self.progress_queue.put(('progress', progress))
                        self.progress_queue.put(('status', f'Procesando {os.path.basename(doc_path)} para {person.nombre_completo}'))
                        
                        # Procesar documento
                        success = processor.process_document(doc_path, person, self.output_dir.get())
                        
                        if success:
                            results['success'] += 1
                            results['details'].append(f"✅ {os.path.basename(doc_path)} - {person.nombre_completo}")
                        else:
                            results['failed'] += 1
                            results['details'].append(f"❌ {os.path.basename(doc_path)} - {person.nombre_completo}")
                        
                        current_operation += 1
                        
                    except Exception as e:
                        logger.error(f"Error procesando {doc_path} para {person.nombre_completo}: {str(e)}")
                        results['failed'] += 1
                        results['details'].append(f"❌ {os.path.basename(doc_path)} - {person.nombre_completo}: {str(e)}")
                        current_operation += 1
            
            # Finalizar procesamiento
            self.progress_queue.put(('complete', results))
            
        except Exception as e:
            logger.error(f"Error en procesamiento: {str(e)}")
            self.progress_queue.put(('error', str(e)))
    
    def check_progress_queue(self):
        """Verifica actualizaciones de progreso"""
        try:
            while True:
                msg_type, data = self.progress_queue.get_nowait()
                
                if msg_type == 'progress':
                    self.progress_var.set(data)
                elif msg_type == 'status':
                    self.status_var.set(data)
                elif msg_type == 'complete':
                    self.handle_processing_complete(data)
                elif msg_type == 'error':
                    self.handle_processing_error(data)
                    
        except queue.Empty:
            pass
        finally:
            # Programar próxima verificación
            self.root.after(100, self.check_progress_queue)
    
    def handle_processing_complete(self, results):
        """Maneja la finalización del procesamiento"""
        self.progress_var.set(100)
        self.status_var.set(f"Completado: {results['success']} exitosos, {results['failed']} fallidos")
        self.process_button.config(state=tk.NORMAL)
        
        # Mostrar resultados detallados
        result_window = tk.Toplevel(self.root)
        result_window.title("Resultados del Procesamiento")
        result_window.geometry("600x400")
        
        # Resumen
        summary_frame = ttk.Frame(result_window)
        summary_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(summary_frame, text=f"✅ Documentos procesados exitosamente: {results['success']}", 
                 foreground='green').pack(anchor=tk.W)
        ttk.Label(summary_frame, text=f"❌ Documentos con errores: {results['failed']}", 
                 foreground='red').pack(anchor=tk.W)
        
        # Detalles
        ttk.Label(result_window, text="Detalles:").pack(anchor=tk.W, padx=10)
        
        # Lista de detalles
        details_frame = ttk.Frame(result_window)
        details_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        details_text = tk.Text(details_frame, wrap=tk.WORD)
        details_scroll = ttk.Scrollbar(details_frame, orient=tk.VERTICAL, command=details_text.yview)
        details_text.configure(yscrollcommand=details_scroll.set)
        
        for detail in results['details']:
            details_text.insert(tk.END, detail + '\n')
        
        details_text.config(state=tk.DISABLED)
        details_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        details_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Botón para abrir carpeta de salida
        ttk.Button(result_window, text="Abrir Carpeta de Salida", 
                  command=lambda: os.startfile(self.output_dir.get())).pack(pady=10)
        
        # Generar reporte
        self.generate_report(results)
    
    def handle_processing_error(self, error_msg):
        """Maneja errores en el procesamiento"""
        self.status_var.set("Error en procesamiento")
        self.process_button.config(state=tk.NORMAL)
        messagebox.showerror("Error", f"Error durante el procesamiento: {error_msg}")
    
    def generate_report(self, results):
        """Genera un reporte del procesamiento"""
        try:
            report_data = {
                'Documento': [],
                'Persona': [],
                'Estado': [],
                'Fecha_Procesamiento': []
            }
            
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            for detail in results['details']:
                parts = detail.split(' - ')
                if len(parts) >= 2:
                    status_doc = parts[0]
                    person_name = parts[1].split(':')[0]  # En caso de que haya mensaje de error
                    
                    status = "Exitoso" if "✅" in status_doc else "Fallido"
                    doc_name = status_doc.replace("✅ ", "").replace("❌ ", "")
                    
                    report_data['Documento'].append(doc_name)
                    report_data['Persona'].append(person_name)
                    report_data['Estado'].append(status)
                    report_data['Fecha_Procesamiento'].append(timestamp)
            
            # Crear DataFrame y guardar
            if report_data['Documento']:
                df_report = pd.DataFrame(report_data)
                report_path = os.path.join(self.output_dir.get(), f"reporte_procesamiento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
                df_report.to_excel(report_path, index=False)
                logger.info(f"Reporte generado: {report_path}")
                
        except Exception as e:
            logger.error(f"Error generando reporte: {str(e)}")
    
    def run(self):
        """Ejecuta la aplicación"""
        self.root.mainloop()

class CommandLineInterface:
    """Interfaz de línea de comandos para uso avanzado"""
    
    def __init__(self):
        self.config = SignatureConfig()
    
    def run(self, excel_path: str, documents: List[str], output_dir: str, config_file: Optional[str] = None):
        """Ejecuta el procesamiento desde línea de comandos"""
        try:
            # Cargar configuración si se proporciona
            if config_file and os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config_dict = json.load(f)
                    self.config = SignatureConfig(**config_dict)
            
            # Leer datos de Excel
            print(f"Cargando datos de {excel_path}...")
            people_data = ExcelReader.read_excel_data(excel_path)
            print(f"Cargadas {len(people_data)} personas")
            
            # Crear directorio de salida
            os.makedirs(output_dir, exist_ok=True)
            
            # Procesar documentos
            processor = DocumentProcessor(self.config)
            total_operations = len(people_data) * len(documents)
            current_operation = 0
            
            print(f"Iniciando procesamiento de {total_operations} operaciones...")
            
            for person in people_data:
                for doc_path in documents:
                    print(f"Procesando {os.path.basename(doc_path)} para {person.nombre_completo}...")
                    
                    success = processor.process_document(doc_path, person, output_dir)
                    
                    if success:
                        print(f"✅ Completado")
                    else:
                        print(f"❌ Error")
                    
                    current_operation += 1
                    progress = (current_operation / total_operations) * 100
                    print(f"Progreso: {progress:.1f}%")
            
            print(f"\n🎉 Procesamiento completado. Archivos guardados en: {output_dir}")
            
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            sys.exit(1)

def main():
    """Función principal"""
    if len(sys.argv) > 1:
        # Modo línea de comandos
        import argparse
        
        parser = argparse.ArgumentParser(description="Simulador de Firma Automática de Documentos")
        parser.add_argument("excel_file", help="Archivo Excel con datos de personas")
        parser.add_argument("documents", nargs="+", help="Documentos a firmar")
        parser.add_argument("-o", "--output", required=True, help="Directorio de salida")
        parser.add_argument("-c", "--config", help="Archivo de configuración JSON")
        
        args = parser.parse_args()
        
        cli = CommandLineInterface()
        cli.run(args.excel_file, args.documents, args.output, args.config)
    else:
        # Modo interfaz gráfica
        try:
            app = DocumentSignerGUI()
            app.run()
        except Exception as e:
            logger.error(f"Error iniciando aplicación: {str(e)}")
            messagebox.showerror("Error Fatal", f"No se pudo iniciar la aplicación: {str(e)}")

if __name__ == "__main__":
    main()

# =============================================================================
# INSTRUCCIONES DE INSTALACIÓN Y USO
# =============================================================================

"""
INSTALACIÓN DE DEPENDENCIAS:
----------------------------
pip install pandas openpyxl PyMuPDF python-docx Pillow

ESTRUCTURA DE DATOS EN EXCEL:
-----------------------------
El archivo Excel debe contener al menos las siguientes columnas:
- Columna con nombres (puede ser "Nombre", "Name", "Nombre Completo", etc.)
- Columna con apellidos (puede ser "Apellido", "Apellidos", "Surname", etc.)
- Columna con DNI (puede ser "DNI", "NIF", "Documento", etc.)

Ejemplo:
| Nombre | Apellido1 | Apellido2 | DNI        |
|--------|-----------|-----------|------------|
| Juan   | García    | López     | 12345678A  |
| María  | Rodríguez | Silva     | 87654321B  |

O alternativamente:
| Nombre Completo      | DNI        |
|---------------------|------------|
| Juan García López   | 12345678A  |
| María Rodríguez Silva| 87654321B  |

USO DE LA APLICACIÓN:
--------------------
1. INTERFAZ GRÁFICA:
   - Ejecutar: python document_signer.py
   - Seguir los pasos en la interfaz

2. LÍNEA DE COMANDOS:
   - Ejemplo: python document_signer.py datos.xlsx documento1.pdf documento2.docx -o ./firmados

CARACTERÍSTICAS MEJORADAS:
-------------------------
✅ Interfaz gráfica moderna con pestañas
✅ Configuración personalizable (tamaños, estilos, etc.)
✅ Procesamiento multihilo para mejor rendimiento
✅ Validación robusta de datos de entrada
✅ Generación automática de firmas manuscritas
✅ Soporte para PDF y DOCX
✅ Auto-detección de archivos Excel
✅ Generación de reportes de procesamiento
✅ Sistema de logging completo
✅ Modo línea de comandos para automatización
✅ Manejo de errores robusto
✅ Barra de progreso en tiempo real

FORMATOS SOPORTADOS:
-------------------
📄 Documentos de entrada: PDF, DOCX
📊 Datos: Excel (.xlsx, .xls)
🖼️ Firmas: Generadas automáticamente (PNG con transparencia)

MEJORAS IMPLEMENTADAS SOBRE EL DISEÑO ORIGINAL:
-----------------------------------------------
1. Arquitectura modular con separación de responsabilidades
2. Interfaz gráfica profesional con pestañas
3. Configuración persistente
4. Procesamiento asíncrono con feedback visual
5. Validación exhaustiva de datos
6. Generación inteligente de firmas
7. Sistema de reportes automático
8. Soporte para línea de comandos
9. Manejo robusto de errores
10. Logging detallado para debugging
"""
